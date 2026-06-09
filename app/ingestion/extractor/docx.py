"""Word(docx) 텍스트 추출 (FR-002) — python-docx 본문 문단 + 표 [Pipeline].

이미지·도형은 제외하고 문단 텍스트와 표 셀 텍스트만 추출한다. 표는 행 단위로
``cell | cell`` 직렬화한다. 라이브러리는 지연 import 한다.
"""

from __future__ import annotations

import io


def extract_docx_text(content: bytes) -> str:
    """docx 바이너리에서 본문 문단·표 텍스트를 추출한다."""
    from docx import Document

    document = Document(io.BytesIO(content))
    parts: list[str] = [
        paragraph.text.strip() for paragraph in document.paragraphs if paragraph.text.strip()
    ]
    for table in document.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                parts.append(" | ".join(cells))
    return "\n".join(parts).strip()
